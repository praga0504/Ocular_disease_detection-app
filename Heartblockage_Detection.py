import os
from datetime import datetime
import streamlit as st
from io import BytesIO
from PIL import Image
from docx import Document
from docx.shared import Inches
import google.generativeai as genai


api_key = "AIzaSyALhvLT1-feVVGK_yN4d8176phhEpp6XEY"  
if api_key is None:
    raise ValueError("GEMINI_API_KEY is not set in environment variables")
genai.configure(api_key=api_key)

generation_config = {
    "temperature": 1,
    "top_p": 0.95,
    "top_k": 64,
    "max_output_tokens": 8192,
    "response_mime_type": "text/plain",
}

model = genai.GenerativeModel(
    model_name="gemini-1.5-pro",
    generation_config=generation_config,
)

# Function to analyze fundus image via API
def analyze_fundus_image_with_api(image_file):
    image = Image.open(image_file)
    current_date = datetime.now().strftime('%Y-%m-%d')

    # Define prompt for Generative AI
    prompt = f"""
Analyze this fundus image or external Image of eye and generate a detailed report. The report should include:
1. Possible signs of heart blockage or related conditions like papilledema, ocular diseas, cataract and related infection.
2. Risk assessment (low, medium, high).
3. Recommendations for the patient.
4. Caution and disclaimer for AI-based analysis.
5. Reporting date: {current_date}.
"""
    chat_session = model.start_chat(history=[])
    response = chat_session.send_message([prompt, image])
    return response.text

# Function to create a downloadable Word report
def create_fundus_report(report_text, fundus_image):
    doc = Document()
    doc.add_heading('Ocular Disease Detection Report', 0)

    for line in report_text.split("\n"):
        if line.strip():
            doc.add_paragraph(line.strip())

    doc.add_heading('Fundus Image:', level=1)
    image_stream = BytesIO(fundus_image.getvalue())
    doc.add_picture(image_stream, width=Inches(6))

    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

# Main Streamlit App
def main():
    st.title("Generative AI-Based Ocular Disease Detection")

    st.header("Upload Fundus Image")
    fundus_image = st.file_uploader("Upload Fundus Image (JPG/PNG)", type=["jpg", "png"])

    if fundus_image:
        st.image(fundus_image, caption='Uploaded Fundus Image', use_column_width=True)

        if st.button("Generate Report"):
            with st.spinner("Analyzing the image..."):
                report = analyze_fundus_image_with_api(fundus_image)

            st.header("Analysis Report")
            st.markdown(report)

            # Allow report download
            doc_file_stream = create_fundus_report(report, fundus_image)
            st.download_button(
                label="Download Report",
                data=doc_file_stream,
                file_name="Fundus_Report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

if __name__ == '__main__':
    main()